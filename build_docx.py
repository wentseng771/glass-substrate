from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

docx_path = r"C:\Users\yehwe\Downloads\wt-ccagent\再痛也沒關係_陳勢安_吉他譜.docx"

content = """陳勢安　－　再痛也沒關係
Key: B　Play: G
Capo 4

臺視偶像劇《花是愛》片尾曲

前奏：
│　G　Am7　│　Bm7　C　│　G　Am　│　Dsus4　D　│
│　Bm7　C　│　D/#F　Bm7　Em　─　│　Am　D　│　Gsus4　G　│

G　　　　　　　　D/#F
　看不見愛情消失的痕跡
Cmaj7　　　　　　G
　聽不見心碎瞬間的聲音
　　　C　　　D　　　Bm7　　Em
找不到一個同情字句　來偽裝這場大雨
　　Am7　　　　　　　D
被淋溼的過去　得不到安息

Cmaj7　　　　　　G
＃ 抽一根沒有溫度的別離
Cmaj7　　　　　　G
　寫一首自以為的悲情
　　　C　　　D　　　Bm7　　Em
等一個預設的結局　好讓我輸到最徹底
　　Am　　　　　　　　D　　　　│D│
我們之間丟下了　一個斷句　無法再繼續

　　　　　G　　　　　　　　D/#F
＠ 我的心被你懸在　到不了的天際
　　Em　　　　　　│　Bm7　─　Dm（G）　│
想念瀰漫著空氣　快不能呼吸
　　　C　　　D　　　　Bm7　　Em
一個人背著幸福練習　擁抱卻沒有力氣
　　　　Am7　　　　　　　C　　D
你穿過　我的身體　回頭卻來　不～及
　　　　　G　　　　　　　　D/#F
我的心被雨困在　揮不去的記憶
　　Em　　　　　　│　Bm7　─　Dm（G）　│
眼淚蒸發了思緒　不讓我看清
　G　　C　　　D　　　　Bm7　　Em
兩個人變成一種或許　等待也顯得多餘
　　　　Am7　　　　　C　　D　（Gsus4）
這份愛　早就已經麻痺　再痛也～沒關係

│　Gsus4　│　G　│

Repeat ＃，＠

間奏：
│　G　│　D/#F　B/D#　│　Em　│　Dm　G　│
│　C　D　│　Bm7　Em　│　C　│　D　│

Repeat ＠

尾奏：
│　G　D/#F　│　Em　D　│
│　C　Bm7　│　Am7　D　│　G　│"""

doc = Document()

# 頁面邊距縮小，讓等寬字型有足夠空間
from docx.shared import Inches
section = doc.sections[0]
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)

for line in content.split('\n'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(line)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    # 標題行加粗
    if '再痛也沒關係' in line and '陳勢安' in line:
        run.bold = True
        run.font.size = Pt(13)

doc.save(docx_path)
print(f'Done: {docx_path}')
